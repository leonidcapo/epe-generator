from __future__ import annotations

import io
import json

from docx import Document

from agents.novelty_checker import candidato_id


def render_candidatos_md(filas: list[dict], warnings: list[str]) -> str:
    lines = ["# Candidatos de semilla EPE", ""]
    if warnings:
        lines += ["> ⚠️ " + w for w in warnings] + [""]
    if not filas:
        lines.append("_No se generaron candidatos._")
        return "\n".join(lines)
    for i, row in enumerate(filas, 1):
        c = row["candidato"]
        score = row["score_llm"]
        score_txt = f"{score:.1f}" if score is not None else "s/valorar (LLM degradado)"
        ajuste_txt = ", ".join(c.covariables_ajuste) if c.covariables_ajuste else "(ninguna)"
        lines += [
            f"## {i}. {c.eje} × {c.subpoblacion} → {c.outcome}",
            "",
            f"**Pregunta tentativa (observacional, multivariado ajustado):** ¿Cuál es la "
            f"asociación de {c.eje} con {c.outcome} en {c.subpoblacion} de la cohorte EPE, "
            f"ajustando por {ajuste_txt}?",
            "",
            f"- Covariables de ajuste: {ajuste_txt}",
            f"- n disponible (conjunto): {c.n_disponible}",
            f"- **Score LLM:** {score_txt}",
            f"- **Score de novedad (0=saturado,1=vacío en literatura):** {row['novedad']:.2f}",
        ]
        just = row["justificacion"].strip()
        if just:
            lines += ["", f"**Justificación:** {just}"]
        lines.append("")
    return "\n".join(lines)


def render_candidatos_json(filas: list[dict]) -> str:
    items = []
    for row in filas:
        c = row["candidato"]
        items.append({
            "id": candidato_id(c),
            "eje": c.eje,
            "subpoblacion": c.subpoblacion,
            "outcome": c.outcome,
            "covariables_ajuste": list(c.covariables_ajuste),
            "n_disponible": c.n_disponible,
            "novedad": row["novedad"],
            "score_llm": row["score_llm"],
        })
    return json.dumps(items, ensure_ascii=False, indent=2)


_PROTO_SECCIONES = {
    "introduccion": "Introducción", "marco_teorico": "Marco Teórico",
    "objetivos": "Objetivos", "hipotesis": "Hipótesis", "metodos": "Métodos",
}


def render_protocolo_md(protocolo) -> str:
    p = protocolo
    lines = [f"# Protocolo — {p.candidato_id}", "", "## PICOT", ""]
    for k, v in p.picot.items():
        lines.append(f"- **{k}:** {v}")
    lines += ["", "## Variables", "", "| nombre | rol | tipo | escala |", "|---|---|---|---|"]
    for v in p.variables:
        lines.append(f"| {v['nombre']} | {v['rol']} | {v['tipo']} | {v.get('escala') or ''} |")
    lines += ["", "## Diseño", "", f"- **tipo:** {p.diseno['tipo']}", f"- **modelo:** {p.diseno['modelo']}"]
    if p.diseno.get("anclajes"):
        lines.append(f"- **anclajes:** {', '.join(p.diseno['anclajes'])}")
    for sec, titulo in _PROTO_SECCIONES.items():
        lines += ["", f"## {titulo}", "", p.prosa.get(sec, "")]
    lines += ["", "## Limitaciones", ""]
    for lim in p.limitaciones:
        lines.append(f"- {lim}")
    if p.warnings_auditoria:
        lines += ["", "> ⚠️ Auditoría:"] + [f"> - {w}" for w in p.warnings_auditoria]
    return "\n".join(lines)


def render_protocolo_docx(protocolo) -> bytes:
    """Word estructurado del protocolo, desde los mismos datos que render_protocolo_md
    (no re-parsea el .md). Devuelve bytes (en memoria)."""
    p = protocolo
    doc = Document()
    doc.add_heading(f"Protocolo — {p.candidato_id}", level=0)

    doc.add_heading("PICOT", level=1)
    for k, v in p.picot.items():
        doc.add_paragraph(f"{k}: {v}", style="List Bullet")

    doc.add_heading("Variables", level=1)
    tabla = doc.add_table(rows=1, cols=4)
    hdr = tabla.rows[0].cells
    hdr[0].text, hdr[1].text, hdr[2].text, hdr[3].text = "nombre", "rol", "tipo", "escala"
    for v in p.variables:
        celdas = tabla.add_row().cells
        celdas[0].text = str(v["nombre"])
        celdas[1].text = str(v["rol"])
        celdas[2].text = str(v["tipo"])
        celdas[3].text = str(v.get("escala") or "")

    doc.add_heading("Diseño", level=1)
    doc.add_paragraph(f"tipo: {p.diseno['tipo']}", style="List Bullet")
    doc.add_paragraph(f"modelo: {p.diseno['modelo']}", style="List Bullet")
    if p.diseno.get("anclajes"):
        doc.add_paragraph(f"anclajes: {', '.join(p.diseno['anclajes'])}", style="List Bullet")

    for sec, titulo in _PROTO_SECCIONES.items():
        doc.add_heading(titulo, level=1)
        doc.add_paragraph(p.prosa.get(sec, ""))

    doc.add_heading("Limitaciones", level=1)
    for lim in p.limitaciones:
        doc.add_paragraph(lim, style="List Bullet")

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
