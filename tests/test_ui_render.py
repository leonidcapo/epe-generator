import io
import json

from docx import Document

from agents.novelty_checker import Candidato
from agents.protocol_designer import Protocolo
from ui_render import render_candidatos_md, render_candidatos_json, render_protocolo_docx, render_protocolo_md


def _fila():
    c = Candidato(eje="riesgo_sistemico_asa", subpoblacion="adultos_mayores",
                 outcome="grado_cooperacion",
                 covariables_ajuste=("farmacoterapia_polifarmacia", "procedencia_acceso"),
                 n_disponible=45)
    return {"candidato": c, "score_llm": 8.5, "justificacion": "relevante", "novedad": 0.9}


def test_render_candidatos_md_incluye_datos_clave():
    md = render_candidatos_md([_fila()], warnings=["aviso x"])
    assert "riesgo_sistemico_asa" in md
    assert "adultos_mayores" in md
    assert "n disponible (conjunto): 45" in md
    assert "farmacoterapia_polifarmacia" in md
    assert "procedencia_acceso" in md
    assert "aviso x" in md


def test_render_candidatos_md_sin_covariables_muestra_ninguna():
    c = Candidato(eje="riesgo_sistemico_asa", subpoblacion="adultos_mayores",
                 outcome="grado_cooperacion", covariables_ajuste=(), n_disponible=45)
    fila = {"candidato": c, "score_llm": None, "justificacion": "", "novedad": 0.5}
    md = render_candidatos_md([fila], warnings=[])
    assert "(ninguna)" in md


def test_render_candidatos_md_vacio():
    md = render_candidatos_md([], warnings=[])
    assert "No se generaron candidatos" in md


def test_render_candidatos_json_roundtrip_campos():
    data = json.loads(render_candidatos_json([_fila()]))
    assert data[0]["eje"] == "riesgo_sistemico_asa"
    assert data[0]["n_disponible"] == 45
    assert data[0]["score_llm"] == 8.5
    assert data[0]["covariables_ajuste"] == ["farmacoterapia_polifarmacia", "procedencia_acceso"]


def _protocolo():
    return Protocolo(
        candidato_id="riesgo_sistemico_asa_asa3_alto_riesgo_nivel_tratamiento_requerido_adj_farmacoterapia_polifarmacia",
        picot={
            "poblacion": "asa3_alto_riesgo", "exposicion": "riesgo_sistemico_asa",
            "covariables_ajuste": "farmacoterapia_polifarmacia",
            "comparador": "categorías de referencia de las covariables",
            "outcome": "nivel_tratamiento_requerido", "tiempo": "transversal (sin seguimiento)",
        },
        variables=[
            {"nombre": "nivel_tratamiento_requerido", "rol": "outcome", "tipo": "categorico", "escala": "ordinal"},
            {"nombre": "riesgo_sistemico_asa", "rol": "exposicion_principal", "tipo": "categorica"},
            {"nombre": "farmacoterapia_polifarmacia", "rol": "covariable", "tipo": "categorica"},
        ],
        diseno={
            "tipo": "transversal_analitico", "modelo": "logistica_ordinal", "anclajes": [],
            "outcome_tipo": "categorico", "outcome_escala": "ordinal",
        },
        prosa={
            "introduccion": "Texto de introducción.", "marco_teorico": "Texto de marco.",
            "objetivos": "Texto de objetivos.", "hipotesis": "Texto de hipótesis.",
            "metodos": "Texto de métodos.",
        },
        limitaciones=["Al ser un estudio observacional transversal, solo permite asociaciones."],
        warnings_auditoria=[],
    )


def test_render_protocolo_md_incluye_secciones_clave():
    md = render_protocolo_md(_protocolo())
    assert "riesgo_sistemico_asa" in md
    assert "logistica_ordinal" in md
    assert "Texto de introducción." in md
    assert "Limitaciones" in md
    assert "Al ser un estudio observacional transversal" in md


def test_render_protocolo_md_sin_anclajes_no_los_menciona():
    md = render_protocolo_md(_protocolo())
    assert "anclajes" not in md.lower()


def test_render_protocolo_docx_produce_bytes_validos():
    data = render_protocolo_docx(_protocolo())
    assert isinstance(data, bytes)
    assert len(data) > 0
    assert data[:2] == b"PK"  # .docx es un archivo zip


def test_render_protocolo_docx_incluye_auditoria_warnings():
    proto = _protocolo()
    # Modificar el protocolo para agregar warnings_auditoria
    proto.warnings_auditoria = ["Lenguaje causal detectado.", "Posible confusión residual."]
    data = render_protocolo_docx(proto)
    doc = Document(io.BytesIO(data))
    # Concatenar todo el texto del documento
    full_text = "\n".join([p.text for p in doc.paragraphs])
    assert "Auditoría" in full_text
    assert "Lenguaje causal detectado." in full_text
    assert "Posible confusión residual." in full_text


def test_render_protocolo_docx_sin_auditoria_warnings_no_incluye_seccion():
    proto = _protocolo()
    proto.warnings_auditoria = []
    data = render_protocolo_docx(proto)
    doc = Document(io.BytesIO(data))
    full_text = "\n".join([p.text for p in doc.paragraphs])
    # La sección "Auditoría" no debe aparecer si warnings_auditoria está vacío
    assert not any("Auditoría" in p.text for p in doc.paragraphs)
